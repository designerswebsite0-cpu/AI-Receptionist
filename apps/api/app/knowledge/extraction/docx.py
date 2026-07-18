from io import BytesIO

from docx import Document

from app.knowledge.extraction.base import ExtractedContent


def extract_docx(content: bytes) -> ExtractedContent:
    document = Document(BytesIO(content))

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    raw_text = "\n".join(parts)
    return ExtractedContent(
        raw_text=raw_text,
        extraction_method="python-docx",
        word_count=len(raw_text.split()),
    )
